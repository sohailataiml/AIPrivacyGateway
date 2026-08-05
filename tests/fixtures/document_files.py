"""Builders for real PDF and DOCX bytes.

Extraction tests need files a real parser will accept, not fixtures shaped to
match the code under test. A hand-written blob that only pypdf's error-recovery
path can read would let a broken extractor pass, so the PDF here is structurally
valid: correct object offsets, a real cross-reference table, and text in a
content stream that any reader can find.

DOCX is built with python-docx itself. That is not circular -- the writer and
the reader are different code paths in the library, and using the writer is how
a test gets a file Word would also open.

The deliberately broken variants at the bottom are the important ones. A parser
is only as good as what it refuses.
"""

from __future__ import annotations

import io
import zipfile
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from collections.abc import Sequence


def make_pdf(pages: Sequence[str]) -> bytes:
    """A structurally valid PDF with one text-bearing page per entry.

    Offsets in the cross-reference table are computed from the bytes actually
    written, so the file is valid rather than merely recoverable.
    """
    if not pages:
        raise ValueError("a PDF needs at least one page")

    page_count = len(pages)
    # Object numbering: 1 catalog, 2 page tree, 3 font, then per page a page
    # object and a content stream.
    font_object = 3
    first_page_object = 4

    objects: dict[int, bytes] = {}
    kids = " ".join(f"{first_page_object + index * 2} 0 R" for index in range(page_count))
    objects[1] = b"<< /Type /Catalog /Pages 2 0 R >>"
    objects[2] = f"<< /Type /Pages /Kids [{kids}] /Count {page_count} >>".encode()
    objects[font_object] = b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>"

    for index, text in enumerate(pages):
        page_object = first_page_object + index * 2
        content_object = page_object + 1
        objects[page_object] = (
            f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            f"/Resources << /Font << /F1 {font_object} 0 R >> >> "
            f"/Contents {content_object} 0 R >>"
        ).encode()
        stream = _content_stream(text)
        objects[content_object] = (
            f"<< /Length {len(stream)} >>\nstream\n".encode() + stream + b"\nendstream"
        )

    out = bytearray(b"%PDF-1.7\n")
    offsets: dict[int, int] = {}
    for number in sorted(objects):
        offsets[number] = len(out)
        out += f"{number} 0 obj\n".encode() + objects[number] + b"\nendobj\n"

    xref_offset = len(out)
    highest = max(objects)
    out += f"xref\n0 {highest + 1}\n".encode()
    out += b"0000000000 65535 f \n"
    for number in range(1, highest + 1):
        out += f"{offsets[number]:010d} 00000 n \n".encode()
    out += f"trailer\n<< /Size {highest + 1} /Root 1 0 R >>\n".encode()
    out += f"startxref\n{xref_offset}\n%%EOF\n".encode()
    return bytes(out)


def _content_stream(text: str) -> bytes:
    """One text-showing operator per line, so line structure survives."""
    lines = text.split("\n") if text else [""]
    parts = [b"BT\n/F1 12 Tf\n14 TL\n72 720 Td\n"]
    for line in lines:
        parts.append(b"(" + _escape(line) + b") Tj\nT*\n")
    parts.append(b"ET")
    return b"".join(parts)


def _escape(text: str) -> bytes:
    """Escape a PDF literal string.

    Backslash first: escaping it after the parentheses would double-escape the
    backslashes this function itself introduced.
    """
    escaped = text.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")
    return escaped.encode("latin-1", "replace")


def make_docx(paragraphs: Sequence[str], table: Sequence[Sequence[str]] = ()) -> bytes:
    """A real DOCX, written by python-docx.

    ``table`` exists because forms put names and identifiers in table cells, and
    an extractor that walks only paragraphs would miss them entirely -- a
    detection gap that looks like an extraction detail.
    """
    from docx import Document as DocxDocument

    document = DocxDocument()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    if table:
        built = document.add_table(rows=len(table), cols=len(table[0]))
        for row_index, row in enumerate(table):
            for cell_index, value in enumerate(row):
                built.cell(row_index, cell_index).text = value

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# Files a parser must refuse
# ---------------------------------------------------------------------------
def make_zip_bomb(*, entries: int = 4, uncompressed_mib: int = 64) -> bytes:
    """A ZIP whose members expand enormously from almost nothing.

    Highly compressible filler, so the central directory honestly declares a
    huge uncompressed size against a tiny compressed one. The guard reads those
    declared sizes and refuses before decompressing anything, which is the only
    order that helps -- a check performed after expansion has already lost.
    """
    filler = b"\0" * (uncompressed_mib * 1024 * 1024)
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED, compresslevel=9) as archive:
        # Named like a DOCX so the file is refused for its shape, not its names.
        archive.writestr("[Content_Types].xml", filler)
        for index in range(entries - 1):
            archive.writestr(f"word/document{index}.xml", filler)
    return buffer.getvalue()


def make_zip_with_many_entries(count: int) -> bytes:
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as archive:
        for index in range(count):
            archive.writestr(f"word/part{index}.xml", b"<x/>")
    return buffer.getvalue()


def make_encrypted_pdf() -> bytes:
    """A PDF with an encryption dictionary.

    Built by pypdf so the flag is set the way a real reader sets it.
    """
    from pypdf import PdfReader, PdfWriter

    writer = PdfWriter()
    reader = PdfReader(io.BytesIO(make_pdf(["locked"])))
    for page in reader.pages:
        writer.add_page(page)
    writer.encrypt("correct-horse-battery-staple")

    buffer = io.BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


TRUNCATED_PDF = b"%PDF-1.7\n1 0 obj\n<< /Type /Catalog"
"""A header and the beginning of one object. Nothing to parse."""

HEADER_ONLY_PDF = b"%PDF-1.7\n%%EOF\n"
"""Passes Phase 1's magic-byte check and contains no pages."""

NOT_A_ZIP_DOCX = b"PK\x03\x04" + b"\xff" * 512
"""Passes Phase 1's DOCX magic-byte check and is not an archive."""


__all__ = [
    "HEADER_ONLY_PDF",
    "NOT_A_ZIP_DOCX",
    "TRUNCATED_PDF",
    "make_docx",
    "make_encrypted_pdf",
    "make_pdf",
    "make_zip_bomb",
    "make_zip_with_many_entries",
]
